import streamlit as st
import requests
from docx import Document
from fpdf import FPDF
import fitz  # PyMuPDF
import os
from io import BytesIO
import unicodedata

# API Setup
API_URL = 'https://openrouter.ai/api/v1/chat/completions'

headers = {
    'Authorization': f'Bearer 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',
    'Content-Type': 'application/json'
}

# Utility functions
def extract_text(file):
    if file.name.endswith(".pdf"):
        doc = fitz.open(stream=file.read(), filetype="pdf")
        return " ".join(page.get_text() for page in doc)
    elif file.name.endswith(".docx"):
        doc = Document(file)
        return "\n".join(p.text for p in doc.paragraphs)
    return ""

def generate_docx(text):
    doc = Document()
    doc.add_paragraph(text)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def sanitize_text(text):
    return unicodedata.normalize("NFKD", text).encode("latin-1", "ignore").decode("latin-1")

def generate_pdf(text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    safe_text = sanitize_text(text)
    for line in safe_text.split("\n"):
        pdf.cell(200, 10, txt=line, ln=True)
    pdf_bytes = pdf.output(dest='S').encode('latin-1')
    buffer = BytesIO(pdf_bytes)
    buffer.seek(0)
    return buffer

# App UI
st.set_page_config(page_title="GotYouCovered", page_icon="📝", layout="centered")
st.title("📝 GotYouCovered")
st.markdown("Generate a customized cover letter using your resume and job description with AI ✨")

st.markdown("---")

# Input Section
col1, col2 = st.columns(2)
with col1:
    resume_file = st.file_uploader("📄 Upload Resume (.pdf or .docx)", type=["pdf", "docx"])
with col2:
    tone = st.selectbox("🎯 Choose a Tone", ["Professional", "Friendly", "Enthusiastic"])

with st.form("cover_letter_form"):
    job_title = st.text_input("💼 Job Title or Description")

    with st.expander("➕ Add Important Points (Optional)"):
        additional_points = st.text_area("💬 Anything specific you'd like to highlight?")

    generate_button = st.form_submit_button("🚀 Generate Cover Letter")

# Session state to hold generated text
if 'final_text' not in st.session_state:
    st.session_state.final_text = None

# Generate cover letter
if generate_button and resume_file and job_title:
    resume_text = extract_text(resume_file)

    prompt = f"""
    Write a {tone.lower()} cover letter for the following job:
    Job: {job_title}
    Resume: {resume_text[:4000]}
    """

    if additional_points.strip():
        prompt += f"\n\nImportant points to include:\n{additional_points.strip()}"

    with st.spinner("Generating your custom cover letter..."):
        data = {
            "model": "deepseek/deepseek-chat:free",
            "messages": [
                {
                    "role": "system",
                    "content": "You are a cover letter generator. Extract the candidate's name, email, and any contact info "
                               "from the resume and include it in the letter. Only return the final cover letter text. "
                               "Do not explain, apologize, or output anything else."
                },
                {"role": "user", "content": prompt}
            ]
        }

        os.environ['NO_PROXY'] = 'openrouter.ai'
        response = requests.post(API_URL, json=data, headers=headers)

        if response.status_code == 200:
            st.session_state.final_text = response.json()["choices"][0]["message"]["content"]
        else:
            st.error(f"Failed to fetch data from API. Status Code: {response.status_code}")

# Display output and download options
if st.session_state.final_text:
    st.markdown("---")
    st.subheader("📄 Generated Cover Letter")
    st.text_area("✍️ Output", st.session_state.final_text, height=300)

    st.markdown("### 📥 Download Cover Letter")
    file_format = st.selectbox("Select Format", ["TXT", "DOCX", "PDF"])
    save_text = st.session_state.final_text

    if file_format == "TXT":
        txt_buffer = BytesIO(save_text.encode("utf-8"))
        st.download_button("📄 TXT", txt_buffer, file_name="Cover_Letter.txt")

    elif file_format == "DOCX":
        file = generate_docx(save_text)
        st.download_button("📝 DOCX", file, file_name="Cover_Letter.docx")

    elif file_format == "PDF":
        file = generate_pdf(save_text)
        st.download_button("📕 PDF", file, file_name="Cover_Letter.pdf")

# Footer
st.markdown("---")
